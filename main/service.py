import os
from datetime import datetime

from django.core.files.base import ContentFile
from docx import Document
from rest_framework.pagination import PageNumberPagination
from rest_framework.response import Response

from main.models import BuyRequest


class PaginationHousings(PageNumberPagination):
    page_size = 10
    max_page_size = 100
    page_size_query_param = 'limit'

    # page_query_param = 'offset'

    def get_paginated_response(self, data):
        return Response({
            # 'links': {
            #     'next': self.get_next_link(),
            #     'previous': self.get_previous_link()
            # },
            'count': self.page.paginator.count,
            'result': data
        })


def format_date(date):
    day = ('0' + str(date.day))[-2:]
    month = ('0' + str(date.month))[-2:]
    return f'{day}.{month}.{date.year}'


def create_contract(pk):
    dir_path = os.path.abspath('./media/shablons/')
    doc = Document(dir_path + '\shablon_dogovor.docx')
    buy_request = BuyRequest.objects.get(pk=pk)
    date_begin = buy_request.date_begin
    date_end = buy_request.date_end
    landlord = buy_request.product.housing.owner
    renter = buy_request.user

    # Замена фамилии в нескольких местах документа
    for paragraph in doc.paragraphs:
        if 'ГОРОД' in paragraph.text:
            paragraph.text = paragraph.text.replace('ГОРОД', buy_request.product.housing.city.name)
        if 'dd' in paragraph.text and 'mm' in paragraph.text and 'yyyy' in paragraph.text:
            now = datetime.now()
            paragraph.text = paragraph.text.replace('dd', ('0' + str(now.day))[-2:])
            paragraph.text = paragraph.text.replace('mm', ('0' + str(now.month))[-2:])
            paragraph.text = paragraph.text.replace('yyyy', str(now.year))
        if 'ФИО_АРЕНДОДАТЕЛЬ' in paragraph.text:
            paragraph.text = paragraph.text.replace('ФИО_АРЕНДОДАТЕЛЬ',
                                                    landlord.get_full_name())
        if 'ФИО_НАНИМАТЕЛЬ' in paragraph.text:
            paragraph.text = paragraph.text.replace('ФИО_НАНИМАТЕЛЬ', renter.get_full_name())
        if 'АДРЕС' in paragraph.text:
            paragraph.text = paragraph.text.replace('АДРЕС', buy_request.product.housing.address)
        if 'ДАТА_НАЧАЛА' in paragraph.text:
            paragraph.text = paragraph.text.replace('ДАТА_НАЧАЛА', format_date(date_begin))
        if 'ДАТА_КОНЦА' in paragraph.text:
            paragraph.text = paragraph.text.replace('ДАТА_КОНЦА', format_date(date_end))
        if 'ЦЕНА' in paragraph.text:
            paragraph.text = paragraph.text.replace('ЦЕНА', str(buy_request.price))
        if 'ИНИЦИАЛЫ_АРЕНДОДАТЕЛЬ' in paragraph.text:
            paragraph.text = paragraph.text.replace('ИНИЦИАЛЫ_АРЕНДОДАТЕЛЬ', landlord.get_initial())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if 'ФАМИЛИЯ' in cell.text:
                    cell.text = cell.text.replace('ФАМИЛИЯ', 'Новая фамилия')
                if 'ПАСПОРТ_СЕРИЯ_АРЕНДОДАТЕЛЬ' in cell.text:
                    cell.text = cell.text.replace('ПАСПОРТ_СЕРИЯ_АРЕНДОДАТЕЛЬ', landlord.passport_series)
                if 'ПАСПОРТ_НОМЕР_АРЕНДОДАТЕЛЬ' in cell.text:
                    cell.text = cell.text.replace('ПАСПОРТ_НОМЕР_АРЕНДОДАТЕЛЬ', landlord.passport_number)
                if 'ПАСПОРТ_ВЫДАН_АРЕНДОДАТЕЛЬ' in cell.text:
                    cell.text = cell.text.replace('ПАСПОРТ_ВЫДАН_АРЕНДОДАТЕЛЬ', landlord.passport_from)
                if 'ПАСПОРТ_ЗАРЕГИСТРИРОВАН_АРЕНДОДАТЕЛЬ' in cell.text:
                    cell.text = cell.text.replace('ПАСПОРТ_ЗАРЕГИСТРИРОВАН_АРЕНДОДАТЕЛЬ',
                                                  landlord.passport_registration_address)
                if 'ИНИЦИАЛЫ_АРЕНДОДАТЕЛЬ' in cell.text:
                    cell.text = cell.text.replace('ИНИЦИАЛЫ_АРЕНДОДАТЕЛЬ', landlord.get_initial())
                if 'ПАСПОРТ_СЕРИЯ_НАНИМАТЕЛЬ' in cell.text:
                    cell.text = cell.text.replace('ПАСПОРТ_СЕРИЯ_НАНИМАТЕЛЬ', renter.passport_series)
                if 'ПАСПОРТ_НОМЕР_НАНИМАТЕЛЬ' in cell.text:
                    cell.text = cell.text.replace('ПАСПОРТ_НОМЕР_НАНИМАТЕЛЬ', renter.passport_number)
                if 'ПАСПОРТ_ВЫДАН_НАНИМАТЕЛЬ' in cell.text:
                    cell.text = cell.text.replace('ПАСПОРТ_ВЫДАН_НАНИМАТЕЛЬ', renter.passport_from)
                if 'ПАСПОРТ_ЗАРЕГИСТРИРОВАН_НАНИМАТЕЛЬ' in cell.text:
                    cell.text = cell.text.replace('ПАСПОРТ_ЗАРЕГИСТРИРОВАН_НАНИМАТЕЛЬ',
                                                  renter.passport_registration_address)
                if 'ИНИЦИАЛЫ_НАНИМАТЕЛЬ' in cell.text:
                    cell.text = cell.text.replace('ИНИЦИАЛЫ_НАНИМАТЕЛЬ', renter.get_initial())

    # Сохранение измененного документа во временный файл
    temp_file_path = dir_path + '//temp.docx'
    doc.save(temp_file_path)
    file_name = f'Dogovor_{datetime.now()}.docx'
    with open(temp_file_path, 'rb') as f:
        buy_request.contract.save(file_name, ContentFile(f.read()), save=True)

    buy_request.save()
    os.remove(temp_file_path)
    return buy_request
